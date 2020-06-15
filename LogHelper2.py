import os
import random
import subprocess
import sys
import time
from datetime import date, timedelta, datetime
from pprint import pprint

import keyboard
import pandas as pd
import win32com.client as win32
from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtCore import Qt
from PyQt5.QtGui import QPalette, QColor
from PyQt5.QtWidgets import QAbstractItemView
from docx import *
from mailmerge import MailMerge
from sqlalchemy import create_engine, exc, inspect, select, MetaData, \
    Column, Integer, String
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker

from Ui_Settings import Ui_Dialog
from ui_Import import Ui_Import
from ui_LogHelper_mainWindow import Ui_MainWindow
from ui_editText import Ui_editText
from ui_search import Ui_Form

date_today = date.strftime(date.today(), '%Y-%m-%d')
daily_handovers = '2'
dpr_prefix = 'DPR-189762-KIRT CHOUEST-'
shf_prefix = 'SHF-189762-KIRT CHOUEST-'
survey_log = []

Base = declarative_base()
engine = create_engine('sqlite:///dpr.db')
Base.metadata.create_all(engine)
Session = sessionmaker(bind=engine)
session = Session()  # Main session
Session2 = sessionmaker(bind=engine)
session2 = Session2()  # Second thread session

# reflect db schema to MetaData
metadata = MetaData()
metadata.reflect(bind=engine)

# use alter statement if table needs added rows
# engine.execute('ALTER TABLE "dpr" ADD dpr_export_msg varchar(255);')

inspector = inspect(engine)
# uncomment to see DB tables
# pprint(inspector.get_table_names())

# uncomment to see DB tables and rows
# pprint(metadata.tables)


def key_maker(bk, num):
    num = str(num)
    if len(num) is 1:
        num = '0' + num
    return bk + num


def random_color():
    r = str(random.randint(0, 255))
    b = str(random.randint(0, 255))
    g = str(random.randint(0, 255))
    rbg = r + ',' + b + ',' + g
    return rbg


def write_pdf(pdf_file_name, docx_path):
    error_message = ''
    try:
        check = subprocess.check_call(
            ('cmd', '/C', 'docto', '-f', docx_path, '-O', pdf_file_name, '-T', 'wdFormatPDF', '-Q'))
    except Exception as e:
        print(e)
        error_message = str(e)
        error_message = '<b><font color="red">PDF file export failed.</font></b><br>PDF Exporter ' + error_message[-35:]
        try:
            check = int(error_message[-4:-1])
        except Exception as ee:
            print(ee)
            check = 999
    print(str(check))
    pdf_success = (check == 0)
    if pdf_success:
        os.startfile(pdf_file_name)
        return '<b><font color="green">PDF file created successfully!</font></b><br>' + pdf_file_name
    elif check == 400:
        return error_message + '<p>MS Word must be installed for PDF output'
    elif check == 220:
        return error_message + '<p>Please make sure that a file with this name is not in use.'
    else:
        return error_message


def folder_checker(string_output):
    string_output = string_output.replace('/', '\\')
    if (string_output[-1:] != "\\") and (string_output[-1:] != "/"):
        return string_output + '/'
    else:
        return string_output


def dpr_number_maker(this_date):
    this_date = this_date.split('-')
    year = this_date[0]
    month = this_date[1]
    day = this_date[2]
    now = date(int(year), int(month), int(day))
    then = date(2019, 8, 4)
    num = str((now - then).days)
    return dpr_prefix + num + "-" + month + '-' + day + '-' + year


def shf_number_maker(this_date, am_pm):
    num = shf_number_calculator(this_date)
    num = str(int(num) - 281)
    if am_pm == 'AM':
        num = str(int(num) - 1)
    this_date = this_date.split('-')
    year = this_date[0]
    month = this_date[1]
    day = this_date[2]
    return shf_prefix + num + "-" + month + '-' + day + '-' + year


def shf_number_calculator(this_date):
    this_date_ = this_date.split('-')
    daily = 2
    value = 380
    now = date(int(this_date_[0]), int(this_date_[1]), int(this_date_[2]))
    then = date(2020, 3, 21)
    new_value = str((now - then).days * daily + value)
    return str(new_value)


def send_email(this_pdf_file, this_dpr_number):
    try:
        outlook = win32.Dispatch('outlook.application')
        mail = outlook.CreateItem(0)
        mail.To = 'survey-mcon-dprs-oii@oceaneering.com;matthew.e.bessette@exxonmobil.com;david.b.pane@exxonmobil.com' \
                  ';ljakse@oceaneering.com '
        mail.Subject = this_dpr_number
        mail.HtmlBody = 'Please see attached for DPR'
        mail.Attachments.Add(this_pdf_file)
        mail.Display(False)
    except Exception as e:
        print(e)
        pass


def set_window_view(is_top):
    this_application = MainWindow()
    if is_top:
        this_application.setWindowFlags(this_application.windowFlags() | QtCore.Qt.WindowStaysOnTopHint)
        this_application.show()
    else:
        this_application.setWindowFlags(this_application.windowFlags() & ~QtCore.Qt.WindowStaysOnTopHint)
        this_application.show()


def create_doc(log_time, log, this_date, make_pdf):
    USER = 1000
    work_complete24 = ''
    pending_work24 = ''
    location_status = ''
    taskReports = ''
    surveyComments = ''
    long_date = this_date.strftime('%B %d, %Y')
    date_selected = this_date.strftime('%Y-%m-%d')
    dpr_key = this_date.strftime('%Y%m%d')
    query = session.query(DPR).get(dpr_key)
    if query:
        work_complete24 = '•    ' + query.dpr_prev24.replace('\n', '\n•    ')
        pending_work24 = '•    ' + query.dpr_next24.replace('\n', '\n•    ')
        location_status = query.dpr_location
        taskReports = query.dpr_task_reports
        surveyComments = query.dpr_comments
    print(dpr_key)
    get_settings = session.query(SettingsTable).get(USER)
    dpr_template = get_settings.dpr_template
    dpr_output_doc = get_settings.dpr_out_doc
    dpr_output_pdf = get_settings.dpr_out_pdf
    name1 = get_settings.name1
    name2 = get_settings.name2
    name3 = get_settings.name3
    title1 = get_settings.title1
    title2 = get_settings.title2
    title3 = get_settings.title3
    shift1 = get_settings.shift1
    shift2 = get_settings.shift2
    shift3 = get_settings.shift3
    handover_value = shf_number_calculator(date_selected)
    dpr_document = MailMerge(dpr_template)
    log_dict = [{'Time1': log_time[0], 'Log1': log[0]}]
    count = 1
    while count < len(log):
        log_dict.append({'Time1': log_time[count], 'Log1': log[count]})
        count += 1
    pprint(log_dict)
    try:
        dpr_document.merge(
            dprNum=dpr_number_maker(date_selected),
            HandoverNum=handover_value,
            Prev24=work_complete24,
            Next24=pending_work24,
            LocStat=location_status,
            SurveyComments=surveyComments,
            TaskReports=taskReports,
            Number=daily_handovers,
            Title1=title1,
            Title2=title2,
            Title3=title3,
            Person1=name1,
            Person2=name2,
            Person3=name3,
            Shift1=shift1,
            Shift2=shift2,
            Shift3=shift3,
            Date=long_date)
        dpr_document.merge_rows('Time1', log_dict)
        dpr_number = dpr_number_maker(date_selected)
        docx_file = folder_checker(dpr_output_doc) + dpr_number + '.docx'
        pdf_file = folder_checker(dpr_output_pdf) + dpr_number + '.pdf'
        dpr_document.write(docx_file)
        if make_pdf:
            time.sleep(0.25)
            print(docx_file)
            print(pdf_file)
            return '<b><font color="green">DOCX file created successfully!</font></b><br>' + docx_file, write_pdf(
                pdf_file, docx_file)
        else:
            return '<b><font color="green">DOCX file created successfully!</font></b><br>' + docx_file, ''
    except (PermissionError, IndexError, FileNotFoundError) as error:
        print(error)
        return '<b><font color="red">DOCX file export failed.</font></b><br>' + error, ''


def create_shf_a(this_date, make_pdf):
    USER = 1000
    tasks_completed = ''
    planned_tasks = ''
    location = ''
    equipment = ''
    comments = ''
    health_safety = ''
    long_date = this_date.strftime('%B %d, %Y')
    date_selected = this_date.strftime('%Y-%m-%d')
    dpr_key = this_date.strftime('%Y%m%d')
    query = session.query(DPR).get(dpr_key)
    if query:
        tasks_completed = '•    ' + query.shf_prev12_a.replace('\n', '\n•    ')
        planned_tasks = '•    ' + query.shf_next12_a.replace('\n', '\n•    ')
        location = query.shf_location_a
        equipment = '•    ' + query.shf_equipment_a
        comments = '•    ' + query.shf_comments_a
        health_safety = '•    ' + query.shf_safety_a
    print(dpr_key)
    get_settings = session.query(SettingsTable).get(USER)
    shf_template = get_settings.shf_template
    shf_output_doc = get_settings.shf_out_doc
    shf_output_pdf = get_settings.shf_out_pdf
    name1 = get_settings.name1
    name2 = get_settings.name2
    shift1 = get_settings.shift1
    shift2 = get_settings.shift2

    try:
        document = MailMerge(shf_template)
        document.merge(
            Location=location,
            TasksCompleted=tasks_completed,
            EquipmentComments=equipment,
            PlannedTasks=planned_tasks,
            HealthSafety=health_safety,
            Comments=comments,
            Person1=name1,
            Person2=name2,
            Shift1=shift1,
            Shift2=shift2,
            Date=long_date)
        shf_number = shf_number_maker(date_selected, 'AM')
        docx_file = folder_checker(shf_output_doc) + shf_number + '.docx'
        pdf_file = folder_checker(shf_output_pdf) + shf_number + '.pdf'
        document.write(docx_file)
        if make_pdf:
            time.sleep(0.25)
            print(docx_file)
            print(pdf_file)
            return '<b><font color="green">DOCX file created successfully!</font></b><br>' + docx_file, write_pdf(
                pdf_file, docx_file)
        else:
            return '<b><font color="green">DOCX file created successfully!</font></b><br>' + docx_file, ''
    except (PermissionError, IndexError, FileNotFoundError) as error:
        print(error)
        return '<b><font color="red">DOCX file export failed.</font></b><br>' + error, ''


def create_shf_b(this_date, make_pdf):
    USER = 1000
    tasks_completed = ''
    planned_tasks = ''
    location = ''
    equipment = ''
    comments = ''
    health_safety = ''
    long_date = this_date.strftime('%B %d, %Y')
    date_selected = this_date.strftime('%Y-%m-%d')
    dpr_key = this_date.strftime('%Y%m%d')
    query = session.query(DPR).get(dpr_key)
    if query:
        tasks_completed = '•    ' + query.shf_prev12_b.replace('\n', '\n•    ')
        planned_tasks = '•    ' + query.shf_next12_b.replace('\n', '\n•    ')
        location = query.shf_location_b
        equipment = '•    ' + query.shf_equipment_b
        comments = '•    ' + query.shf_comments_b
        health_safety = '•    ' + query.shf_safety_b
    print(dpr_key)
    get_settings = session.query(SettingsTable).get(USER)
    shf_template = get_settings.shf_template
    shf_output_doc = get_settings.shf_out_doc
    shf_output_pdf = get_settings.shf_out_pdf
    name1 = get_settings.name2
    name2 = get_settings.name1
    shift1 = get_settings.shift2
    shift2 = get_settings.shift1

    try:
        shf_document = MailMerge(shf_template)
        shf_document.merge(
            Location=location,
            TasksCompleted=tasks_completed,
            EquipmentComments=equipment,
            PlannedTasks=planned_tasks,
            HealthSafety=health_safety,
            Comments=comments,
            Person1=name1,
            Person2=name2,
            Shift1=shift1,
            Shift2=shift2,
            Date=long_date)
        shf_number = shf_number_maker(date_selected, 'PM')
        docx_file = folder_checker(shf_output_doc) + shf_number + '.docx'
        pdf_file = folder_checker(shf_output_pdf) + shf_number + '.pdf'
        shf_document.write(docx_file)
        if make_pdf:
            time.sleep(0.25)
            print(docx_file)
            print(pdf_file)
            return '<b><font color="green">DOCX file created successfully!</font></b><br>' + docx_file, write_pdf(
                pdf_file, docx_file)
        else:
            return '<b><font color="green">DOCX file created successfully!</font></b><br>' + docx_file, ''
    except (PermissionError, IndexError, FileNotFoundError) as error:
        print(error)
        return '<b><font color="red">DOCX file export failed.</font></b><br>' + error, ''


class Log(Base):
    __tablename__ = 'log'
    id = Column(Integer, primary_key=True)
    line_number = Column(Integer, nullable=False)
    time = Column(String, nullable=False)
    date = Column(String, nullable=False)
    comment = Column(String, nullable=False)
    note = Column(String, nullable=True)


class SettingsTable(Base):
    __tablename__ = 'settings'
    id = Column(Integer, primary_key=True)
    dpr_template = Column(String, nullable=False)
    dpr_out_pdf = Column(String, nullable=False)
    dpr_out_doc = Column(String, nullable=False)
    shf_template = Column(String, nullable=False)
    shf_out_pdf = Column(String, nullable=True)
    shf_out_doc = Column(String, nullable=True)
    dpr_file_prefix = Column(String, nullable=True)
    pdf_file_prefix = Column(String, nullable=True)
    name1 = Column(String, nullable=True)
    name2 = Column(String, nullable=True)
    name3 = Column(String, nullable=True)
    title1 = Column(String, nullable=True)
    title2 = Column(String, nullable=True)
    title3 = Column(String, nullable=True)
    shift1 = Column(String, nullable=True)
    shift2 = Column(String, nullable=True)
    shift3 = Column(String, nullable=True)


class DPR(Base):
    __tablename__ = 'dpr'
    id = Column(Integer, primary_key=True)
    dpr_prev24 = Column(String, nullable=True)
    dpr_next24 = Column(String, nullable=True)
    dpr_date = Column(String, nullable=True)
    dpr_location = Column(String, nullable=True)
    dpr_task_reports = Column(String, nullable=True)
    dpr_comments = Column(String, nullable=True)
    shf_prev12_a = Column(String, nullable=True)
    shf_next12_a = Column(String, nullable=True)
    shf_location_a = Column(String, nullable=True)
    shf_equipment_a = Column(String, nullable=True)
    shf_safety_a = Column(String, nullable=True)
    shf_comments_a = Column(String, nullable=True)
    shf_prev12_b = Column(String, nullable=True)
    shf_next12_b = Column(String, nullable=True)
    shf_location_b = Column(String, nullable=True)
    shf_equipment_b = Column(String, nullable=True)
    shf_safety_b = Column(String, nullable=True)
    shf_comments_b = Column(String, nullable=True)
    dpr_export_msg = Column(String, nullable=True)
    shf_export_msg_a = Column(String, nullable=True)
    shf_export_msg_b = Column(String, nullable=True)


class LoadFiles(QtCore.QThread):
    progress_tracker = QtCore.pyqtSignal(int)
    data_ready = QtCore.pyqtSignal(pd.DataFrame)
    files = []

    def set_files(self, files):
        self.files = files

    def run(self):
        list_all = []
        progress_count = 0
        for file in self.files:
            if len(file) > 0:
                this_date = file[-15:-5].split('-')
                corrected_date = this_date[2] + '-' + this_date[0] + '-' + this_date[1]
                base_key = this_date[2] + this_date[0] + this_date[1]
                print(corrected_date, base_key)
                dpr_document = Document(file)
                log_table = dpr_document.tables[6]
                log_list = []
                count = 0
                line_number = 0
                for rows in log_table.rows:
                    if count > 1 and len(rows.cells[0].text) > 0:
                        log_list.append([key_maker(base_key, line_number),
                                         corrected_date,
                                         rows.cells[0].text,
                                         rows.cells[1].text,
                                         line_number])
                        line_number += 1
                    count += 1
                list_all.extend(log_list)
            progress_count += 1
            self.progress_tracker.emit(progress_count)
        df = pd.DataFrame(list_all, columns=['key', 'date', 'time', 'comment', 'line_number'])
        final_df = df.sort_values(by=['key'], ascending=True)
        self.data_ready.emit(final_df)


class SaveToDatabase(QtCore.QThread):
    progress_tracker = QtCore.pyqtSignal(int)
    job_complete = QtCore.pyqtSignal(int)
    log_df = pd.DataFrame()

    def set_df(self, df):
        self.log_df = df

    def run(self):
        print('IMPORTING DATA!')
        log_table = Log()
        session2.rollback()
        count = 0
        for index, row in self.log_df.iterrows():
            count += 1
            self.progress_tracker.emit(count)
            try:
                log_table.id = int(row['key'])
                log_table.date = row['date']
                log_table.time = row['time']
                log_table.comment = row['comment']
                log_table.line_number = row['line_number']
                session2.add(log_table)
                log_table = Log()
                session2.commit()
            except exc.IntegrityError as error:
                print(error)
                session2.rollback()
                update = session2.query(Log).get(row['key'])
                update.date = row['date']
                update.time = row['time']
                update.comment = row['comment']
                update.line_number = row['line_number']
                session2.commit()
                continue
        self.job_complete.emit(count)


class ImportWindow(QtWidgets.QWidget, Ui_Import):
    def __init__(self):
        super(ImportWindow, self).__init__()
        Ui_Import.__init__(self)
        self.setupUi(self)
        self.log = Log()
        self.results_table.setEditTriggers(QAbstractItemView.NoEditTriggers)  # EDITS OFF
        self.results_table.setColumnWidth(0, 70)
        self.results_table.setColumnWidth(1, 45)
        self.results_table.setColumnWidth(2, 340)
        self.final_df = pd.DataFrame()
        self.b_import_data.setEnabled(False)
        self.b_import_data.clicked.connect(self._save_to_db)
        self.b_select_files.clicked.connect(self._open_files_dialog)
        self.progressBar.setHidden(True)
        self.results_string = ''
        self.results_table.setStyleSheet('QTableWidget::item:selected{ background-color: rgb(126, 208, 228); }')
        self.layoutWidget.setObjectName('OnlyMe')
        self.setStyleSheet('QWidget#OnlyMe {background-color: rgb(255, 255, 255);}')

    def _save_to_db(self):
        self.results_string = 'Updating database...'
        self.progressBar.setHidden(False)
        self.save_db = SaveToDatabase()
        self.save_db.set_df(self.final_df)
        self.progressBar.setMaximum(len(self.final_df.index))
        self.save_db.job_complete.connect(self.job_complete)
        self.save_db.progress_tracker.connect(self.update_progress_bar)
        self.save_db.start()

    def job_complete(self, value):
        self.results_label.setText(str(value) + ' rows saved to database.')
        self.progressBar.setValue(0)
        self.progressBar.setHidden(True)
        self.b_import_data.setEnabled(False)

    def _open_files_dialog(self):
        options = QtWidgets.QFileDialog.Options()
        QtWidgets.QFileSystemModel(self)
        # noinspection PyCallByClass
        files, _ = QtWidgets.QFileDialog.getOpenFileNames(self, "Select Files to Load", "", "DOCX Files (*.docx)",
                                                          options=options)
        self.results_string = 'Loading data...'
        self.progressBar.setHidden(False)
        self.load_files = LoadFiles()
        self.load_files.set_files(files)
        self.progressBar.setMaximum(len(files))
        self.load_files.progress_tracker.connect(self.update_progress_bar)
        self.load_files.data_ready.connect(self.get_df)
        self.load_files.start()
        self._reload()

    def get_df(self, df):
        self._load_all_data(df)
        self.final_df = df
        self.progressBar.setValue(0)
        self.progressBar.setHidden(True)
        self.b_import_data.setEnabled(True)

    def update_progress_bar(self, value):
        self.results_label.setText(self.results_string)
        self.progressBar.setValue(value)

    def _extract_data_from_docs(self):
        pass

    def _load_all_data(self, final_df):
        try:
            row_number = 0
            last_date = ''
            color_bool = False
            for index, row in final_df.iterrows():
                self.results_table.insertRow(row_number)
                self.results_table.setRowHeight(row_number, 2.5)
                self.results_table.setItem(row_number, 0, QtWidgets.QTableWidgetItem())
                self.results_table.setItem(row_number, 1, QtWidgets.QTableWidgetItem())
                self.results_table.setItem(row_number, 2, QtWidgets.QTableWidgetItem())
                date_item = self.results_table.item(row_number, 0)
                time_item = self.results_table.item(row_number, 1)
                comment_item = self.results_table.item(row_number, 2)
                date_item.setTextAlignment(QtCore.Qt.AlignCenter)
                time_item.setTextAlignment(QtCore.Qt.AlignCenter)
                date_item.setText(row['date'])
                time_item.setText(row['time'])
                comment_item.setText(row['comment'])
                if last_date != date_item.text():
                    color_bool = not color_bool
                if color_bool:
                    date_item.setBackground(QtGui.QColor(208, 236, 249))
                    time_item.setBackground(QtGui.QColor(208, 236, 249))
                    comment_item.setBackground(QtGui.QColor(208, 236, 249))
                print(row_number)
                self.results_label.setText(str(row_number + 1) + ' lines imported')
                row_number += 1
                last_date = date_item.text()
        except (exc.InvalidRequestError, Exception) as e:
            print(e)

    def _reload(self):
        self.results_table.setRowCount(0)


class SearchWindow(QtWidgets.QWidget, Ui_Form):
    def __init__(self):
        super(SearchWindow, self).__init__()
        Ui_Form.__init__(self)
        self.setupUi(self)
        self.log = Log()
        self._load_all_data()
        self.b_refresh.pressed.connect(self._reload)
        self.results_table.setEditTriggers(QAbstractItemView.NoEditTriggers)  # EDITS OFF
        self.search_input.textChanged.connect(self._reload)
        self.results_table.setColumnWidth(0, 70)
        self.results_table.setColumnWidth(1, 45)
        self.results_table.setColumnWidth(2, 340)
        self.results_label.setText('')

    def _load_all_data(self):
        search_string = self.search_input.text()
        try:
            if search_string == '':
                select_this = engine.execute('SELECT * FROM "log"')
            else:
                search_string = '%' + search_string + '%'
                select_this = engine.execute('SELECT * FROM "log" WHERE comment LIKE "' + search_string + '"')
            row_number = 0
            last_date = ''
            color_bool = False
            for data in select_this:
                self.results_table.insertRow(row_number)
                self.results_table.setRowHeight(row_number, 2.5)
                self.results_table.setItem(row_number, 0, QtWidgets.QTableWidgetItem())
                self.results_table.setItem(row_number, 1, QtWidgets.QTableWidgetItem())
                self.results_table.setItem(row_number, 2, QtWidgets.QTableWidgetItem())
                date_item = self.results_table.item(row_number, 0)
                time_item = self.results_table.item(row_number, 1)
                comment_item = self.results_table.item(row_number, 2)
                date_item.setTextAlignment(QtCore.Qt.AlignCenter)
                time_item.setTextAlignment(QtCore.Qt.AlignCenter)
                date_item.setText(data.date)
                time_item.setText(data.time)
                comment_item.setText(data.comment)
                if last_date != date_item.text():
                    color_bool = not color_bool
                if color_bool:
                    date_item.setBackground(QtGui.QColor(208, 236, 249))
                    time_item.setBackground(QtGui.QColor(208, 236, 249))
                    comment_item.setBackground(QtGui.QColor(208, 236, 249))
                print(row_number)
                row_number += 1
                last_date = date_item.text()
        except (exc.InvalidRequestError, Exception) as e:
            print(e)

    def _reload(self):
        self.results_table.setRowCount(0)
        self._load_all_data()


class EditText(QtWidgets.QMainWindow, Ui_editText):
    def __init__(self, key, type_message):
        self.key = key
        pprint(type_message)
        self.export_flag = False
        print(key)
        print(type_message)
        self.flag = True
        super(EditText, self).__init__()
        self.setupUi(self)
        self.setWindowFlags(self.windowFlags() | Qt.FramelessWindowHint | Qt.WindowStaysOnTopHint)
        if len(key) is 10:
            self.obj = Log()
        else:
            self.obj = DPR()
            self.message = type_message[1].replace('\n', '<br>')
            self.report_type = int(type_message[0])
            self.export_flag = True
        self.setAttribute(Qt.WA_TranslucentBackground, True)
        if self.export_flag:
            self.setStyleSheet("background-color : rgba(195,195,195,210); border:0;")
        else:
            self._change_bg()
        self.textEdit.setStyleSheet("background-color : rgba(255,255,255,200); border:0;")
        self.label.setStyleSheet("background-color : rgba(255,255,255,0); border:0;")
        self.closeButton.setStyleSheet("background-color : rgba(255,255,255,0); border:0;")
        self.shortcut_close = QtWidgets.QShortcut(QtGui.QKeySequence("Ctrl+n"), self)
        self.shortcut_close.activated.connect(self.close_window)
        self.change_bg = QtWidgets.QShortcut(QtGui.QKeySequence("Ctrl+shift+b"), self)
        self.change_bg.activated.connect(self._change_bg)
        font_label = QtGui.QFont()
        font_label.setPointSize(18)
        self.label.setFont(font_label)
        self.label.setContentsMargins(0, 0, 5, 0)
        if self.export_flag:
            self.update = session.query(DPR).get(key)
            font_edit_text = QtGui.QFont()
            font_edit_text.setPointSize(10)
            self.setFixedSize(500, 225)
            self.textEdit.setFont(font_edit_text)
            self.textEdit.setHtml(self.message)
            self.label.setText('EXPORT RESULTS')
            self.closeButton.pressed.connect(self.close_window)
            self.textEdit.setEnabled(False)
        else:
            self.update = session.query(Log).get(key)
            self.textEdit.setHtml(self.update.note)
            self.label.setText(self.update.date + ' - ' + self.update.time)
            self.closeButton.pressed.connect(self.close_window)

        self._drag_active = False
        self.previous_pos = None
        self._setup()
        self.show()

    def _setup(self):
        pass

    def _change_bg(self):
        self.setStyleSheet("background-color : rgba(" + random_color() + ",190); border:0;")

    def mousePressEvent(self, e):
        self.previous_pos = e.globalPos()

    def keyPressEvent(self, keyEvent):
        if self.flag:
            keyboard.press_and_release('tab, ctrl+end')
            self.flag = False

    def mouseMoveEvent(self, e):
        if self.previous_pos is not None:
            delta = e.globalPos() - self.previous_pos
            self.move(self.x() + delta.x(), self.y() + delta.y())
            self.previous_pos = e.globalPos()
        self._drag_active = True

    def mouseReleaseEvent(self, e):
        if self._drag_active:
            self._drag_active = False

    def close_window(self):
        try:
            if self.export_flag:
                if self.report_type == 1:
                    self.update.dpr_export_msg = self.textEdit.toPlainText()
                elif self.report_type == 2:
                    self.update.shf_export_msg_a = self.textEdit.toPlainText()
                elif self.report_type == 3:
                    self.update.shf_export_msg_b = self.textEdit.toPlainText()
            else:
                self.update.note = self.textEdit.toPlainText()
            session.commit()
        except Exception as e:
            print(e)
        self.close()


# noinspection PyCallByClass
class Settings(QtWidgets.QDialog, Ui_Dialog):
    def __init__(self, obj=None):
        super(Settings, self).__init__()
        Ui_Dialog.__init__(self)
        self.USER = 1000  # change if creating multiple profiles
        if obj:
            self.obj = obj
        else:
            self.obj = SettingsTable()

        self.ui = Ui_Dialog()
        self.ui.setupUi(self)
        settings = session.query(SettingsTable).get(self.USER)
        if settings:
            self._load_settings()
        else:
            self.obj.id = self.USER
            self.obj.dpr_template = 'templates/DPR_TEMPLATE.docx'
            self.obj.dpr_out_doc = 'output/'
            self.obj.dpr_out_pdf = 'output/'
            self.obj.shf_template = 'templates/SHF_TEMPLATE.docx'
            self.obj.shf_out_doc = 'output/'
            self.obj.shf_out_pdf = 'output/'
            self.obj.dpr_file_prefix = 'DPR-189762-KIRT CHOUEST-'
            self.obj.shf_file_prefix = 'SHF-189762-KIRT CHOUEST-'
            self.obj.name1 = ''
            self.obj.name2 = ''
            self.obj.name3 = ''
            self.obj.title1 = ''
            self.obj.title2 = ''
            self.obj.title3 = ''
            self.obj.shift1 = ''
            self.obj.shift2 = ''
            self.obj.shift3 = ''
            session.add(self.obj)
            print('Defaults Loaded')
            session.commit()
            self._load_settings()
        self.ui.b_settings_cancel.pressed.connect(self.close)
        self.ui.b_settings_save.pressed.connect(self._save_settings)
        self.ui.b_dpr_template.pressed.connect(lambda: self._open_file_name_dialog(self.ui.s_dpr_template))
        self.ui.b_shf_template.pressed.connect(lambda: self._open_file_name_dialog(self.ui.s_shf_template))
        self.ui.b_dpr_output_doc.pressed.connect(lambda: self._open_folder_name_dialog(self.ui.s_dpr_output_doc))
        self.ui.b_dpr_output_pdf.pressed.connect(lambda: self._open_folder_name_dialog(self.ui.s_dpr_output_pdf))
        self.ui.b_shf_output_pdf.pressed.connect(lambda: self._open_folder_name_dialog(self.ui.s_shf_output_pdf))
        self.ui.b_shf_output_doc.pressed.connect(lambda: self._open_folder_name_dialog(self.ui.s_shf_output_doc))

    def _open_file_name_dialog(self, s_line_edit):
        options = QtWidgets.QFileDialog.Options()
        file_name, _ = QtWidgets.QFileDialog.getOpenFileName(self, "Select Template File", "", "DOCX Files (*.docx)",
                                                             options=options)
        if file_name:
            print(file_name)
            s_line_edit.setText(file_name)

    # noinspection PyCallByClass
    def _open_folder_name_dialog(self, s_line_edit):
        folder_name = QtWidgets.QFileDialog.getExistingDirectory(self, "Select Output Folder")
        if folder_name:
            print(folder_name)
            s_line_edit.setText(folder_name)

    def _save_settings(self):
        session.rollback()
        update_settings = session.query(SettingsTable).get(self.USER)

        if update_settings:
            update_settings.dpr_template = self.ui.s_dpr_template.text()
            update_settings.dpr_out_doc = self.ui.s_dpr_output_doc.text()
            update_settings.dpr_out_pdf = self.ui.s_dpr_output_pdf.text()
            update_settings.shf_template = self.ui.s_shf_template.text()
            update_settings.shf_out_doc = self.ui.s_shf_output_doc.text()
            update_settings.shf_out_pdf = self.ui.s_shf_output_pdf.text()
            update_settings.name1 = self.ui.s_name1.text()
            update_settings.name2 = self.ui.s_name2.text()
            update_settings.name3 = self.ui.s_name3.text()
            update_settings.title1 = self.ui.s_title1.text()
            update_settings.title2 = self.ui.s_title2.text()
            update_settings.title3 = self.ui.s_title3.text()
            update_settings.shift1 = self.ui.s_shift1.text()
            update_settings.shift2 = self.ui.s_shift2.text()
            update_settings.shift3 = self.ui.s_shift3.text()
            session.commit()
            print('SETTINGS UPDATED')

        else:
            self.obj.id = self.USER
            self.obj.dpr_template = self.ui.s_dpr_template.text()
            self.obj.dpr_out_doc = self.ui.s_dpr_output_doc.text()
            self.obj.dpr_out_pdf = self.ui.s_dpr_output_pdf.text()
            self.obj.shf_template = self.ui.s_shf_template.text()
            self.obj.shf_out_doc = self.ui.s_shf_output_doc.text()
            self.obj.shf_out_pdf = self.ui.s_shf_output_pdf.text()
            session.add(self.obj)
            print('SETTINGS INSERT')
            session.commit()
        self.ui.statusbar.setText('Saved')
        self.ui.b_settings_cancel.setText('Close')

    def _load_settings(self):
        get_settings = session.query(SettingsTable).get(self.USER)
        self.ui.s_dpr_template.setText(get_settings.dpr_template)
        self.ui.s_dpr_output_doc.setText(get_settings.dpr_out_doc)
        self.ui.s_dpr_output_pdf.setText(get_settings.dpr_out_pdf)
        self.ui.s_shf_template.setText(get_settings.shf_template)
        self.ui.s_shf_output_doc.setText(get_settings.shf_out_doc)
        self.ui.s_shf_output_pdf.setText(get_settings.shf_out_pdf)
        self.ui.s_name1.setText(get_settings.name1)
        self.ui.s_name2.setText(get_settings.name2)
        self.ui.s_name3.setText(get_settings.name3)
        self.ui.s_title1.setText(get_settings.title1)
        self.ui.s_title2.setText(get_settings.title2)
        self.ui.s_title3.setText(get_settings.title3)
        self.ui.s_shift1.setText(get_settings.shift1)
        self.ui.s_shift2.setText(get_settings.shift2)
        self.ui.s_shift3.setText(get_settings.shift3)


def _add_shadow_effect(item):
    effect = QtWidgets.QGraphicsDropShadowEffect()
    effect.setBlurRadius(3)
    effect.setColor(QtGui.QColor(41, 136, 156))
    effect.setOffset(1, 1)
    item.setGraphicsEffect(effect)


def get_lighter(color_num):
    color_num = int(color_num)
    color_num = color_num + 75
    if color_num > 255:
        color_num = 255
    return color_num


def _format_time(time_string):
    if time_string == 't':
        return datetime.now().strftime("%H:%M")
    time_string = time_string.replace(':', '').replace(' ', '')
    if len(time_string) is 0:
        return time_string
    if len(time_string) is 1:
        time_string = '000' + time_string
    if len(time_string) is 2:
        time_string = '00' + time_string
    if len(time_string) is 3:
        time_string = '0' + time_string
    time_string = time_string[0:2] + ':' + time_string[2:]
    return time_string


class MainWindow(QtWidgets.QMainWindow, Ui_MainWindow):
    def __init__(self, obj=None, dpr=None, settings=None):
        super(MainWindow, self).__init__()
        Ui_MainWindow.__init__(self)
        if obj:
            self.obj = obj
        else:
            self.obj = Log()
        if dpr:
            self.dpr = dpr
        else:
            self.dpr = DPR()
        if settings:
            self.settings = settings
        else:
            self.settings = SettingsTable()
        self.ui = Ui_MainWindow()
        self.ui.setupUi(self)
        self.log = []
        self.log_time = []
        self.copy_mode = False
        self._setup()
        self.sort_log_bool = self.ui.actionAuto_Sort_Log.isChecked()
        self._load_db()
        self.ui.actionAuto_Sort_Log.triggered.connect(self._sort_log_switch)
        self.changes_made_toggle = False
        self.shortcut = QtWidgets.QShortcut(QtGui.QKeySequence("Ctrl+s"), self)
        self.shortcut.activated.connect(self._save_to_db)
        self.shortcut_next_day = QtWidgets.QShortcut(QtGui.QKeySequence("Ctrl+Shift+Right"), self)
        self.shortcut_next_day.activated.connect(self._date_next)
        self.shortcut_prev_day = QtWidgets.QShortcut(QtGui.QKeySequence("Ctrl+Shift+Left"), self)
        self.shortcut_prev_day.activated.connect(self._date_previous)
        self.shortcut_notes = QtWidgets.QShortcut(QtGui.QKeySequence("Ctrl+n"), self)
        self.shortcut_notes.activated.connect(lambda: self._notes(self.item))
        self.shortcut_random_bg = QtWidgets.QShortcut(QtGui.QKeySequence("Ctrl+Shift+g"), self)
        self.shortcut_random_bg.activated.connect(self._random_bg)
        self.item = self.ui.tableWidget.currentCellChanged.connect(self._get_item)
        self.ui.tableWidget.currentCellChanged.connect(self._save_to_db)
        self.ui.b_save.pressed.connect(self._create_doc)
        self.setMinimumSize(352, 504)
        self.setBaseSize(352, 504)
        self.resize(352, 504)
        self.ui.dateEdit.dateChanged.connect(self._load_db)
        self.ui.b_date_next.pressed.connect(self._date_next)
        self.ui.b_date_previous.pressed.connect(self._date_previous)
        self.ui.actionSettings.triggered.connect(self._show_settings)
        self.ui.dpr_prev24.selectionChanged.connect(self._save_dpr)
        self.ui.dpr_next24.selectionChanged.connect(self._save_dpr)
        self.ui.dpr_location.selectionChanged.connect(self._save_dpr)
        self.ui.dpr_task_reports.selectionChanged.connect(self._save_dpr)
        self.ui.dpr_comments.selectionChanged.connect(self._save_dpr)
        self.ui.b_clear_dpr.pressed.connect(self._clear_dpr_form)
        self.ui.b_save_dpr.pressed.connect(self._save_dpr)
        self.ui.b_save_shf_a.pressed.connect(self._save_shf_a)
        self.ui.b_save_shf_b.pressed.connect(self._save_shf_b)
        self.ui.dpr_prev24.setTabChangesFocus(True)
        self.ui.dpr_next24.setTabChangesFocus(True)
        self.ui.dateEdit.dateChanged.connect(self._load_forms)
        self.ui.actionCopy_Day.triggered.connect(self._toggle_copy_mode)
        self.ui.actionPaste_Day.triggered.connect(self._paste_all)
        self.ui.actionRandom_Background.triggered.connect(self._random_bg)
        self.ui.actionSave.triggered.connect(self._paste_all)
        palette1 = QPalette()
        palette1.setColor(QPalette.Window, QColor(142, 210, 225))
        palette1.setColor(QPalette.Base, QColor(232, 232, 232))
        palette1.setColor(QPalette.AlternateBase, QColor(143, 186, 209))
        palette2 = QPalette()
        palette2.setColor(QPalette.Window, QColor(224, 245, 252))
        palette2.setColor(QPalette.Base, QColor(232, 232, 232))
        palette2.setColor(QPalette.AlternateBase, QColor(143, 186, 209))
        palette3 = QPalette()
        palette3.setColor(QPalette.Window, QColor(207, 254, 240))
        palette3.setColor(QPalette.Base, QColor(232, 232, 232))
        palette3.setColor(QPalette.AlternateBase, QColor(143, 186, 209))
        palette4 = QPalette()
        palette4.setColor(QPalette.Window, QColor(255, 232, 217))
        palette4.setColor(QPalette.Base, QColor(232, 232, 232))
        palette4.setColor(QPalette.AlternateBase, QColor(143, 186, 209))
        self.ui.tableWidget.setStyleSheet('QTableWidget::item:selected{ background-color: rgb(126, 208, 228); }')
        self.ui.centralwidget.setPalette(palette1)
        self.ui.centralwidget.setAutoFillBackground(True)
        self.ui.tab_2.setPalette(palette2)
        self.ui.tab_3.setPalette(palette2)
        self.ui.tab_4.setPalette(palette3)
        self.ui.tab_5.setPalette(palette4)
        self.ui.tab.setPalette(palette2)
        self.ui.tab.setAutoFillBackground(True)
        self.ui.tab_2.setAutoFillBackground(True)
        self.ui.tab_3.setAutoFillBackground(True)
        self.ui.tab_4.setAutoFillBackground(True)
        self.ui.tab_5.setAutoFillBackground(True)
        self.ui.dpr_prev24.setPalette(QPalette(Qt.white))
        self.ui.dpr_next24.setPalette(QPalette(Qt.white))
        self.ui.dpr_location.setPalette(QPalette(Qt.white))
        self.ui.dpr_task_reports.setPalette(QPalette(Qt.white))
        self.ui.dpr_comments.setPalette(QPalette(Qt.white))
        self.ui.shf_prev12_a.setPalette(QPalette(Qt.white))
        self.ui.shf_prev12_b.setPalette(QPalette(Qt.white))
        self.ui.shf_next12_b.setPalette(QPalette(Qt.white))
        self.ui.shf_next12_a.setPalette(QPalette(Qt.white))
        _add_shadow_effect(self.ui.l_show_date)
        _add_shadow_effect(self.ui.b_date_next)
        _add_shadow_effect(self.ui.b_date_previous)
        _add_shadow_effect(self.ui.dateEdit)
        self.setWindowIcon(QtGui.QIcon("favicon.ico"))
        self._set_shf_shifts()
        self.ui.b_shf_create_a.pressed.connect(self._create_shf_a)
        self.ui.b_shf_create_b.pressed.connect(self._create_shf_b)
        self.ui.b_clear_shf_a.pressed.connect(self._clear_shf_a)
        self.ui.b_clear_shf_b.pressed.connect(self._clear_shf_b)
        self._times_default_theme()
        self.ui.actionWindow_always_in_front.setChecked(False)
        self.ui.actionWindow_always_in_front.triggered.connect(self._change_window_view)
        self.ui.actionWindow_always_in_front.setEnabled(False)
        self.ui.actionClear_Log.triggered.connect(self._menu_clear_log)
        self.ui.actionSearch.triggered.connect(self._open_search_window)
        self.ui.actionImport_Log_Data.triggered.connect(self._open_import_window)
        self.ui.tabWidget.setCurrentIndex(0)

    def _sort_log_switch(self):
        self.sort_log_bool = not self.sort_log_bool
        if self.sort_log_bool:
            if self.ui.tableWidget.currentColumn() is 1:
                self.ui.tableWidget.setCurrentCell(self.ui.tableWidget.currentRow(), 0)
            self._load_db()

    def _open_import_window(self):
        self.dialog = ImportWindow()
        self.dialog.setWindowIcon(QtGui.QIcon('favicon.ico'))
        self.dialog.setWindowFlags(self.dialog.windowFlags() | QtCore.Qt.WindowStaysOnTopHint)
        self.dialog.show()

    def _open_search_window(self):
        self.dialog = SearchWindow()
        self.dialog.setWindowIcon(QtGui.QIcon('favicon.ico'))
        self.dialog.setWindowFlags(self.dialog.windowFlags() | QtCore.Qt.WindowStaysOnTopHint)
        self.dialog.show()

    def _times_default_theme(self):
        for row in range(self.ui.tableWidget.rowCount()):
            item = self.ui.tableWidget.item(row, 0)
            item.setTextAlignment(QtCore.Qt.AlignCenter)
            if row % 2 == 0:
                item.setBackground(QtGui.QColor(208, 236, 249))
            else:
                item.setBackground(QtGui.QColor(216, 239, 250))
            self.ui.tableWidget.setRowHeight(row, 1)
            self._comment_color()

    def _change_window_view(self):
        print(self.ui.actionWindow_always_in_front.isChecked())
        if self.actionWindow_always_in_front.isChecked():
            new_window = MainWindow()
            new_window.setWindowFlags(new_window.windowFlags() | QtCore.Qt.WindowStaysOnTopHint)
        else:
            new_window = MainWindow()
            new_window.setWindowFlags(new_window.windowFlags() & ~QtCore.Qt.WindowStaysOnTopHint)

    def _menu_clear_log(self):
        self._clear_table()
        self.changes_made_toggle = True
        self._save_to_db()

    def _status_bar_update(self, display_string):
        self.ui.statusbar.showMessage(display_string)

    def _toggle_copy_mode(self):
        if self.copy_mode is False:
            self.copy_mode = True
            self.ui.statusbar.showMessage('COPY MODE ENABLED')
            self.ui.actionPaste_Day.setEnabled(True)
            old_date = self.ui.l_show_date.text()
            self.ui.l_show_date.setText('[COPY] - ' + old_date)
            self._copy_color()
        else:
            self.copy_mode = False
            self.ui.statusbar.showMessage('COPY MODE DISABLED', 4000)
            self.ui.actionPaste_Day.setEnabled(False)
            self._load_db()

    def _random_bg(self):
        palette1 = QPalette()
        color1 = random_color().split(',')
        color2 = random_color().split(',')
        color3 = random_color().split(',')
        palette1.setColor(QPalette.Window, QColor(int(color1[0]), int(color1[1]), int(color1[2])))
        palette2 = QPalette()
        palette3 = QPalette()
        palette4 = QPalette()
        palette2.setColor(QPalette.Window,
                          QColor(get_lighter(color1[0]), get_lighter(color1[1]), get_lighter(color1[2])))
        palette3.setColor(QPalette.Window,
                          QColor(get_lighter(color2[0]), get_lighter(color2[1]), get_lighter(color2[2])))
        palette4.setColor(QPalette.Window,
                          QColor(get_lighter(color3[0]), get_lighter(color3[1]), get_lighter(color3[2])))
        self.ui.centralwidget.setPalette(palette1)

        self.ui.tab_2.setPalette(palette2)
        self.ui.tab_4.setPalette(palette3)
        self.ui.tab_5.setPalette(palette4)
        self.ui.tab.setPalette(palette2)

    def _save_dpr(self):
        if self.copy_mode:
            self.copy_mode = False
            self._paste_all()
        else:
            session.rollback()
            dpr_key = str(self.ui.dateEdit.text()).replace('-', '')
            update = session.query(DPR).get(dpr_key)
            print(self.ui.dpr_prev24.toPlainText())
            print(type(self.ui.dpr_prev24.toPlainText()))
            if update:
                update.dpr_prev24 = self.ui.dpr_prev24.toPlainText()
                update.dpr_next24 = self.ui.dpr_next24.toPlainText()
                update.dpr_location = self.ui.dpr_location.text()
                update.dpr_task_reports = self.ui.dpr_task_reports.text()
                update.dpr_comments = self.ui.dpr_comments.text()
                update.dpr_export_msg = ''
                session.commit()
                print('DPR UPDATED')
            else:
                self.dpr.id = int(dpr_key)
                self.dpr.dpr_prev24 = self.ui.dpr_prev24.toPlainText()
                self.dpr.dpr_next24 = self.ui.dpr_next24.toPlainText()
                self.dpr.dpr_location = self.ui.dpr_location.text()
                self.dpr.dpr_task_reports = self.ui.dpr_task_reports.text()
                self.dpr.dpr_date = self.ui.dpr_comments.text()
                self.dpr.dpr_export_msg = ''
                session.add(self.dpr)
                self.dpr = DPR()
                print(dpr_key)
                print('DPR INSERT')
                session.commit()

    def _save_shf_a(self):
        if self.copy_mode:
            self.copy_mode = False
            self._paste_all()
        else:
            self._set_shf_shifts()
            session.rollback()
            dpr_key = str(self.ui.dateEdit.text()).replace('-', '')
            update = session.query(DPR).get(dpr_key)
            if update:
                update.shf_prev12_a = self.ui.shf_prev12_a.toPlainText()
                update.shf_next12_a = self.ui.shf_next12_a.toPlainText()
                update.shf_location_a = self.ui.shf_location_a.text()
                update.shf_equipment_a = self.ui.shf_equipment_a.text()
                update.shf_safety_a = self.ui.shf_safety_a.text()
                update.shf_comments_a = self.ui.shf_comments_a.text()
                session.commit()
                print('SHF A UPDATED')
            else:
                self.dpr.id = int(dpr_key)
                self.dpr.shf_prev12_a = self.ui.shf_prev12_a.toPlainText()
                self.dpr.shf_next12_a = self.ui.shf_next12_a.toPlainText()
                self.dpr.shf_location_a = self.ui.shf_location_a.text()
                self.dpr.shf_equipment_a = self.ui.shf_equipment_a.text()
                self.dpr.shf_safety_a = self.ui.shf_safety_a.text()
                self.dpr.shf_comments_a = self.ui.shf_comments_a.text()
                session.add(self.dpr)
                self.dpr = DPR()
                print(dpr_key)
                print('SHF A INSERT')
                session.commit()

    def _paste_all(self):
        if self.copy_mode:
            self._save_to_db()
            self._sort_times()
            self._save_dpr()
            self._save_shf_a()
            self._save_shf_b()
            self.ui.statusbar.showMessage('PASTED ALL DATA, COPY MODE DISABLED', 4000)
            self.copy_mode = False
            self.ui.actionPaste_Day.setEnabled(False)
            self._times_default_theme()
        else:
            self._save_to_db()
            self._save_dpr()
            self._save_shf_a()
            self._save_shf_b()
            self.ui.statusbar.showMessage('SAVED ALL DATA', 4000)

    def _save_shf_b(self):
        if self.copy_mode:
            self.copy_mode = False
            self._paste_all()
        else:
            self._set_shf_shifts()
            session.rollback()
            dpr_key = str(self.ui.dateEdit.text()).replace('-', '')
            update = session.query(DPR).get(dpr_key)
            if update:
                update.shf_prev12_b = self.ui.shf_prev12_b.toPlainText()
                update.shf_next12_b = self.ui.shf_next12_b.toPlainText()
                update.shf_location_b = self.ui.shf_location_b.text()
                update.shf_equipment_b = self.ui.shf_equipment_b.text()
                update.shf_safety_b = self.ui.shf_safety_b.text()
                update.shf_comments_b = self.ui.shf_comments_b.text()
                session.commit()
                print('SHF B UPDATED')
            else:
                self.dpr.id = int(dpr_key)
                self.dpr.shf_prev12_b = self.ui.shf_prev12_b.toPlainText()
                self.dpr.shf_next12_b = self.ui.shf_next12_b.toPlainText()
                self.dpr.shf_location_b = self.ui.shf_location_b.text()
                self.dpr.shf_equipment_b = self.ui.shf_equipment_b.text()
                self.dpr.shf_safety_b = self.ui.shf_safety_b.text()
                self.dpr.shf_comments_b = self.ui.shf_comments_b.text()
                session.add(self.dpr)
                self.dpr = DPR()
                print(dpr_key)
                print('SHF B INSERT')
                session.commit()

    def _clear_dpr_form(self):
        self.ui.dpr_prev24.setPlainText('')
        self.ui.dpr_next24.setPlainText('')
        self.ui.dpr_location.setText('')
        self.ui.dpr_task_reports.setText('')
        self.ui.dpr_comments.setText('')

    def _clear_shf_a(self):
        self.ui.shf_prev12_a.setPlainText('')
        self.ui.shf_next12_a.setPlainText('')
        self.ui.shf_location_a.setText('')
        self.ui.shf_equipment_a.setText('')
        self.ui.shf_safety_a.setText('')
        self.ui.shf_comments_a.setText('')

    def _clear_shf_b(self):
        self.ui.shf_prev12_b.setPlainText('')
        self.ui.shf_next12_b.setPlainText('')
        self.ui.shf_location_b.setText('')
        self.ui.shf_equipment_b.setText('')
        self.ui.shf_safety_b.setText('')
        self.ui.shf_comments_b.setText('')

    def _load_forms(self):
        if not self.copy_mode:
            self._clear_dpr_form()
            self._clear_shf_a()
            self._clear_shf_b()
            dpr_key = int(self.ui.dateEdit.text().replace('-', ''))
            query = session.query(DPR).get(dpr_key)
            if query:
                self.ui.dpr_prev24.setPlainText(query.dpr_prev24)
                self.ui.dpr_next24.setPlainText(query.dpr_next24)
                self.ui.dpr_location.setText(query.dpr_location)
                self.ui.dpr_task_reports.setText(query.dpr_task_reports)
                self.ui.dpr_comments.setText(query.dpr_comments)
                self.ui.shf_prev12_a.setPlainText(query.shf_prev12_a)
                self.ui.shf_next12_a.setPlainText(query.shf_next12_a)
                self.ui.shf_location_a.setText(query.shf_location_a)
                self.ui.shf_equipment_a.setText(query.shf_equipment_a)
                self.ui.shf_comments_a.setText(query.shf_comments_a)
                self.ui.shf_prev12_b.setPlainText(query.shf_prev12_b)
                self.ui.shf_next12_b.setPlainText(query.shf_next12_b)
                self.ui.shf_location_b.setText(query.shf_location_b)
                self.ui.shf_equipment_b.setText(query.shf_equipment_b)
                self.ui.shf_safety_b.setText(query.shf_safety_b)
                self.ui.shf_safety_a.setText(query.shf_safety_a)
                self.ui.shf_comments_b.setText(query.shf_comments_b)

    def _create_doc(self):
        make_pdf = self.ui.dpr_pdf_checkbox.isChecked()
        if self.ui.tableWidget.currentColumn() == 1:
            self._sort_times()
        self._load_db()
        if len(self.log_time) > 0:
            this_date = self.ui.dateEdit.date().toPyDate()
            try:
                docx_msg, pdf_msg = create_doc(self.log_time, self.log, this_date, make_pdf)
                msg = docx_msg + '\n----------------\n' + pdf_msg
            except Exception as e:
                msg = e
        else:
            msg = 'ERROR - NO LOG DATA FOUND'
        print(msg)
        self._show_export_details(msg, '1')

    def _create_shf_a(self):
        create_pdf = self.ui.b_shf_pdf_a.isChecked()
        self._save_shf_a()
        dpr_key = int(self.ui.dateEdit.text().replace('-', ''))
        query = session.query(DPR).get(dpr_key)
        if query:
            this_date = self.ui.dateEdit.date().toPyDate()
            try:
                docx_msg, pdf_msg = create_shf_a(this_date, create_pdf)
                msg = docx_msg + '\n----------------\n' + pdf_msg
            except Exception as e:
                msg = e
        else:
            msg = 'ERROR - SHF DATA FOUND FOR THIS DATE'
        print(msg)
        self._show_export_details(msg, '2')

    def _create_shf_b(self):
        create_pdf = self.ui.b_shf_pdf_b.isChecked()
        self._save_shf_b()
        dpr_key = int(self.ui.dateEdit.text().replace('-', ''))
        query = session.query(DPR).get(dpr_key)
        if query:
            this_date = self.ui.dateEdit.date().toPyDate()
            try:
                docx_msg, pdf_msg = create_shf_b(this_date, create_pdf)
                msg = docx_msg + '\n----------------\n' + pdf_msg
            except Exception as e:
                msg = e
        else:
            msg = 'ERROR - SHF DATA FOUND FOR THIS DATE'
        print(msg)
        self._show_export_details(msg, '3')

    def _set_shf_shifts(self):
        users = session.query(SettingsTable).get(1000)
        if users:
            self.ui.l_shf_name_a.setText(users.name1)
            self.ui.l_shf_name_b.setText(users.name2)
            self.ui.l_shf_shift_a.setText(users.shift1)
            self.ui.l_shf_shift_b.setText(users.shift2)

    def _setup(self):
        self.ui.dateEdit.setDate(date.today())
        font = QtGui.QFont()
        font.setBold(True)
        font.setPixelSize(13)
        self.ui.tabWidget.setTabText(0, 'LOG')
        self.ui.tabWidget.setTabText(1, 'DPR')
        header_time = QtWidgets.QTableWidgetItem('Time')
        header_comments = QtWidgets.QTableWidgetItem('Comments')
        header_comments.setBackground(QtGui.QColor(231, 232, 225))
        header_comments.setFont(font)
        header_time.setFont(font)
        header_time.setBackground(QtGui.QColor(231, 232, 225))
        # self.ui.tableWidget.setColumnWidth(1, 500)
        self.ui.tableWidget.setColumnWidth(0, 60)
        self.ui.tableWidget.setAlternatingRowColors(True)
        self.ui.tableWidget.setHorizontalHeaderItem(0, header_time)
        self.ui.tableWidget.setHorizontalHeaderItem(1, header_comments)
        self.ui.tableWidget.setColumnCount(2)
        self.ui.tableWidget.setRowCount(50)
        self.ui.tableWidget.viewport().installEventFilter(self)
        self._load_forms()
        self.ui.tableWidget.resizeColumnsToContents()
        self.set_table_width()
        rows = self.ui.tableWidget.rowCount()
        cols = self.ui.tableWidget.columnCount()
        for row in range(rows):
            for col in range(cols):
                self.ui.tableWidget.setItem(row, col, QtWidgets.QTableWidgetItem())

    def set_table_width(self):
        width = self.width()
        self.ui.tableWidget.setColumnWidth(1, width - 86)

    def resizeEvent(self, event):
        self.set_table_width()
        super(MainWindow, self).resizeEvent(event)

    def _date_previous(self):
        new_date = self.ui.dateEdit.text().split('-')
        year = int(new_date[0])
        month = int(new_date[1])
        day = int(new_date[2])
        print(year, day, month)
        self.ui.dateEdit.setDate(date(year, month, day) - timedelta(days=1))

    def _get_item(self):
        row = self.ui.tableWidget.currentRow()
        col = self.ui.tableWidget.currentColumn()
        self.item = self.ui.tableWidget.item(row, col)
        return self.item

    def _show_settings(self):
        self.dialog = Settings()
        self.dialog.setWindowFlags(self.dialog.windowFlags() | QtCore.Qt.WindowStaysOnTopHint)
        self.dialog.show()
        if self.dialog.exec_() == 0:
            self._set_shf_shifts()

    def _sort_times(self):
        this_date = str(self.ui.dateEdit.text())
        try:
            times_list = []
            select_date = engine.execute('SELECT * FROM "log" WHERE date="' + this_date + '"')
            for data in select_date:
                times_list.append([data.id, data.time, data.comment, data.line_number, data.note])
            df = pd.DataFrame(times_list, columns=['id', 'time', 'comment', 'line_number', 'note'])
            final_df = df.sort_values(by=['time'], ascending=True).reset_index()
            existing_data = session.query(Log).all()
            for data in existing_data:
                if data.date == this_date:
                    session.delete(data)
            session.commit()
            count = 0
            for index, row in final_df.iterrows():
                if count < 10:
                    num = '0' + str(count)
                else:
                    num = str(count)
                key = str(self.ui.dateEdit.text()).replace('-', '') + num
                if row['time'] != '':
                    self.obj.id = int(key)
                    self.obj.time = row['time']
                    self.obj.comment = row['comment']
                    self.obj.date = this_date
                    self.obj.line_number = count
                    self.obj.note = row['note']
                    session.add(self.obj)
                    self.obj = Log()
                    count += 1
            session.commit()
        except exc.InvalidRequestError:
            print('exc.InvalidRequestError:')

    def _date_next(self):
        new_date = self.ui.dateEdit.text().split('-')
        year = int(new_date[0])
        month = int(new_date[1])
        day = int(new_date[2])
        print(year, day, month)
        self.ui.dateEdit.setDate(date(year, month, day) + timedelta(days=1))

    def eventFilter(self, obj, event):
        if event.type() == QtCore.QEvent.MouseButtonPress:
            item = self.ui.tableWidget.itemAt(event.pos())
            if event.button() == QtCore.Qt.RightButton:
                self._generate_menu(event.pos(), item)
                # print(item.row(), item.column(), item.text())
        return QtCore.QObject.event(obj, event)

    def _generate_menu(self, pos, item):
        self.ui.tableWidget.setCurrentCell(item.row(), item.column())
        if item.text() is not '':
            notes = QtWidgets.QAction('&Notes', self)
            notes.setShortcut('Ctrl+n')
            notes.setStatusTip('Notes')
            notes.triggered.connect(lambda: self._notes(item))
            menu = QtWidgets.QMenu(self)
            right_click_items = [notes]  # Add more to this list
            menu.addActions(right_click_items)
            menu.exec_(self.ui.tableWidget.mapToGlobal(pos))  # +++

    def _notes(self, item):
        if str(type(item)) == '<class \'PyQt5.QtWidgets.QTableWidgetItem\'>':
            if item.text() is not '':
                if item.row() < 10:
                    num = '0' + str(item.row())
                else:
                    num = str(item.row())
                key = str(self.ui.dateEdit.text()).replace('-', '') + num
                self.dialog = EditText(key, ['', ''])
                self.dialog.setPalette(palette)
                self.dialog.show()

    def _show_export_details(self, msg, rp):
        report_type = str(rp)
        key = str(self.ui.dateEdit.text()).replace('-', '')
        type_message = [report_type, msg]
        self.dialog = EditText(key, type_message)
        self.dialog.setPalette(palette)
        self.dialog.show()

    def _table_changed(self):
        this_date = self.ui.dateEdit.text()
        select_date = engine.execute('SELECT * FROM "log" WHERE date="' + this_date + '"')
        for data in select_date:
            if self.ui.tableWidget.item(data.line_number, 0).text() != data.time:
                return True, data.line_number
            if self.ui.tableWidget.item(data.line_number, 1).text() != data.comment:
                return True, data.line_number
        return False, None

    def _comment_color(self):
        self.ui.tableWidget.setAlternatingRowColors(True)
        for row in range(self.ui.tableWidget.rowCount()):
            comment_item = self.ui.tableWidget.item(row, 1)
            time_item = self.ui.tableWidget.item(row, 0)
            if row % 2 == 0:
                time_item.setBackground(QtGui.QColor(208, 236, 249))
            else:
                time_item.setBackground(QtGui.QColor(216, 239, 250))
            if len(comment_item.text()) > 0:
                if 'rov' in comment_item.text().lower():
                    comment_item.setBackground(QtGui.QColor(251, 225, 253))
                elif 'vessel' in comment_item.text().lower():
                    comment_item.setBackground(QtGui.QColor(222, 254, 207))
                elif 'fix' in comment_item.text().lower():
                    comment_item.setBackground(QtGui.QColor(242, 252, 177))
                elif 'deployed' in comment_item.text().lower() or \
                        'seabed' in comment_item.text().lower():
                    comment_item.setBackground(QtGui.QColor(168, 209, 255))
                else:
                    comment_item.setBackground(QtGui.QColor(255, 255, 249))
            else:
                if row % 2 == 0:
                    comment_item.setBackground(QtGui.QColor(245, 245, 245))
                else:
                    comment_item.setBackground(QtGui.QColor(255, 255, 255))

    def _copy_color(self):
        for row in range(self.ui.tableWidget.rowCount()):
            comment_item = self.ui.tableWidget.item(row, 1)
            time_item = self.ui.tableWidget.item(row, 0)
            comment_item.setBackground(QtGui.QColor(255, 242, 0))
            time_item.setBackground(QtGui.QColor(255, 242, 0))

    def _save_to_db(self):
        if self.copy_mode:
            self.copy_mode = False
            self.changes_made_toggle = True
            self._paste_all()
        else:
            changes_made, line_changed = self._table_changed()
            session.rollback()
            try:
                rows = self.ui.tableWidget.rowCount()
                for row in range(rows):
                    if row < 10:
                        num = '0' + str(row)
                    else:
                        num = str(row)
                    key = str(self.ui.dateEdit.text()).replace('-', '') + num
                    table = metadata.tables['log']
                    select_statement = select([table.c.id]).where(table.c.id == key)
                    search = engine.execute(select_statement)
                    string_result = str(search.fetchall())[2:-3]
                    this_time = _format_time(self.ui.tableWidget.item(row, 0).text())
                    this_comment = self.ui.tableWidget.item(row, 1).text()
                    key_exists = len(string_result) > 0
                    if self.changes_made_toggle:
                        line_changed = row
                        changes_made = True
                    if string_result == key and changes_made and row == line_changed:
                        update = session.query(Log).get(key)
                        update.time = this_time
                        update.comment = this_comment
                        changes_made = True
                        print(' ---------------------------------')
                        print('| -UPDATED  ROW WITH ID: ' + key + ' |')
                        print(' ---------------------------------')

                    elif (len(self.ui.tableWidget.item(row, 0).text()) > 0 or not self.sort_log_bool) and not key_exists:
                        self.obj.id = int(key)
                        self.obj.time = this_time
                        self.obj.comment = this_comment
                        self.obj.date = str(self.ui.dateEdit.text())
                        self.obj.line_number = row
                        session.add(self.obj)
                        self.obj = Log()
                        changes_made = True
                        print(' ---------------------------------')
                        print('| -INSERTED ROW WITH ID: ' + key + ' |')
                        print(' ---------------------------------')
                if changes_made:
                    session.commit()
                    self.ui.statusbar.showMessage('Saved')
                self._load_db()
            except (exc.IntegrityError, exc.OperationalError, exc.InvalidRequestError) as error:
                print(error)
            self.changes_made_toggle = False

    def _clear_table(self):
        rows = self.ui.tableWidget.rowCount()
        cols = self.ui.tableWidget.columnCount()
        for row in range(rows):
            for col in range(cols):
                self.ui.tableWidget.item(row, col).setText('')

    def _load_db(self):
        if not self.copy_mode:
            self._clear_table()
            self.ui.tableWidget.setAlternatingRowColors(True)
            if self.ui.tableWidget.currentColumn() is not 1 and self.sort_log_bool:
                self._sort_times()
            this_date = str(self.ui.dateEdit.text())
            try:
                font_bold = QtGui.QFont()
                font_bold.setBold(True)
                font_bold.setItalic(True)
                font = QtGui.QFont()
                select_date = engine.execute('SELECT * FROM "log" WHERE date="' + this_date + '"')
                self.log = []
                self.log_time = []
                for data in select_date:
                    time_item = self.ui.tableWidget.item(data.line_number, 0)
                    comment_item = self.ui.tableWidget.item(data.line_number, 1)
                    self.log.append(data.comment)
                    self.log_time.append(data.time)
                    time_item.setText(data.time)
                    comment_item.setText(data.comment)
                    if data.note:
                        time_item.setFont(font_bold)
                        comment_item.setFont(font_bold)
                    else:
                        time_item.setFont(font)
                        comment_item.setFont(font)
            except (exc.InvalidRequestError, Exception) as e:
                print(e)
            long_date = self.ui.dateEdit.date().toPyDate()
            self.ui.l_show_date.setText(long_date.strftime('%A - %B %d, %Y'))
            self._comment_color()

    def _clear_db(self):
        existing_data = session.query(Log).all()
        for data in existing_data:
            session.delete(data)
        session.commit()
        self._load_db()
        print('cleared')


palette = QPalette()
palette.setColor(QPalette.Window, QColor(0, 213, 220, 50))
palette.setColor(QPalette.WindowText, QColor(197, 233, 254, 125))
app = QtWidgets.QApplication([])
application = MainWindow()
application.setWindowFlags(application.windowFlags() | QtCore.Qt.WindowStaysOnTopHint)
application.show()
sys.exit(app.exec())
